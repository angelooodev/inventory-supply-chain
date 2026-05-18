import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


def clear_paragraph(paragraph):
    p = paragraph._element
    for child in list(p):
        p.remove(child)


def set_cell_text(cell, value, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = str(value or "")
    for paragraph in cell.paragraphs:
        paragraph.alignment = align


def format_currency(value):
    amount = float(value or 0)
    return f"P{amount:,.2f}"


def insert_signature(paragraph, image_path, width_inches=2.2):
    clear_paragraph(paragraph)
    if image_path and Path(image_path).exists():
        run = paragraph.add_run()
        run.add_picture(str(image_path), width=Inches(width_inches))


def insert_dual_signatures(paragraph, left_image, right_image):
    clear_paragraph(paragraph)
    left_run = paragraph.add_run()
    if left_image and Path(left_image).exists():
        left_run.add_picture(str(left_image), width=Inches(2.0))

    spacer_run = paragraph.add_run(" " * 26)

    right_run = paragraph.add_run()
    if right_image and Path(right_image).exists():
        right_run.add_picture(str(right_image), width=Inches(2.0))

    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT


def main():
    if len(sys.argv) != 4:
        raise SystemExit("Usage: generate_purchase_order.py <template> <payload> <output>")

    template_path = Path(sys.argv[1])
    payload_path = Path(sys.argv[2])
    output_path = Path(sys.argv[3])

    payload = json.loads(payload_path.read_text(encoding="utf-8"))
    doc = Document(str(template_path))

    po_details = doc.tables[0]
    set_cell_text(po_details.cell(1, 0), payload["po_number"])
    set_cell_text(po_details.cell(1, 1), payload["date_issued"])
    set_cell_text(po_details.cell(1, 2), payload["delivery_date"])

    supplier_table = doc.tables[1]
    supplier_values = [
        payload["supplier_name"],
        payload["contact_person"],
        payload["supplier_address"],
        payload["supplier_contact"],
    ]
    for row_index, value in enumerate(supplier_values):
        set_cell_text(supplier_table.cell(row_index, 1), value)

    items_table = doc.tables[2]
    for row_index in range(1, len(items_table.rows)):
        row = items_table.rows[row_index]
        item = payload["items"][row_index - 1] if row_index - 1 < len(payload["items"]) else None
        if item:
            set_cell_text(row.cells[0], item["description"])
            set_cell_text(row.cells[1], item["quantity"], WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[2], item["unit_price"], WD_ALIGN_PARAGRAPH.RIGHT)
            set_cell_text(row.cells[3], item["unit"], WD_ALIGN_PARAGRAPH.CENTER)
            set_cell_text(row.cells[4], item["amount"], WD_ALIGN_PARAGRAPH.RIGHT)
        else:
            for cell in row.cells:
                set_cell_text(cell, "")
    items_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    summary_table = doc.tables[3]
    summary_values = [
        payload["subtotal"],
        payload["vat"],
        payload["shipping_fee"],
        payload["total_amount"],
    ]
    for row_index, value in enumerate(summary_values):
        set_cell_text(summary_table.cell(row_index, 1), value, WD_ALIGN_PARAGRAPH.RIGHT)

    doc.paragraphs[16].text = f"Delivery Address:\n{payload['warehouse_name']}\n{payload['warehouse_address']}"
    doc.paragraphs[34].text = "Prepared By:                                                                                                           Supplier Confirmation:"
    doc.paragraphs[36].text = "Warehouse Manager                                                                                           Supplier Representative"
    doc.paragraphs[37].text = "Approved By:"
    doc.paragraphs[39].text = "Lumiere Corporation CEO"

    insert_dual_signatures(
        doc.paragraphs[35],
        payload.get("warehouse_manager_signature_path", ""),
        payload.get("supplier_signature_path", ""),
    )
    insert_signature(doc.paragraphs[38], payload.get("owner_signature_path", ""), width_inches=2.1)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))


if __name__ == "__main__":
    main()
